"""
Garmin Chat - Standalone Desktop Application
A local desktop chatbot for querying Garmin Connect data.
"""

import sys
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox
import threading
import json
from pathlib import Path
from garmin_handler import GarminDataHandler
from ai_client import AIClient
import logging
from datetime import datetime

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class SettingsDialog(tk.Toplevel):
    """Dialog for managing application settings including AI provider selection"""
    
    def __init__(self, parent, current_config=None, colors=None):
        super().__init__(parent)
        self.title("Settings")
        self.geometry("700x700")
        self.resizable(False, False)
        
        # Set window icon (same as main window)
        try:
            # Get the correct base path for PyInstaller exe
            if getattr(sys, 'frozen', False):
                # Running as compiled executable
                base_path = Path(sys._MEIPASS)
            else:
                # Running as script
                base_path = Path(__file__).parent
            
            icon_path = base_path / "logo.ico"
            if icon_path.exists():
                self.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"Could not load Settings dialog icon: {e}")
        
        # Store colors (use parent's colors or defaults)
        self.colors = colors or {
            'bg': '#F3F3F3',
            'card_bg': '#FFFFFF',
            'text': '#1F1F1F',
            'text_secondary': '#605E5C',
            'border': '#EDEBE9',
            'accent': '#0078D4'
        }
        
        # Apply theme to dialog
        self.configure(bg=self.colors['bg'])
        
        # Make modal
        self.transient(parent)
        self.grab_set()
        
        # Center on parent
        self.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (self.winfo_width() // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (self.winfo_height() // 2)
        self.geometry(f"+{x}+{y}")
        
        self.result = None
        self.current_config = current_config or {}
        
        # Get AI provider info
        from ai_client import AIClient
        self.providers = AIClient.get_available_providers()
        
        # Create StringVars for ALL providers upfront (so they persist when switching)
        # This ensures all keys are saved, not just the currently selected provider
        for provider_id in self.providers.keys():
            if provider_id == 'azure':
                # Azure needs special handling
                if not hasattr(self, 'azure_endpoint_var'):
                    self.azure_endpoint_var = tk.StringVar(value=current_config.get('azure_endpoint', ''))
                if not hasattr(self, 'azure_deployment_var'):
                    self.azure_deployment_var = tk.StringVar(value=current_config.get('azure_deployment', ''))
                if not hasattr(self, 'azure_key_var'):
                    self.azure_key_var = tk.StringVar(value=current_config.get('azure_api_key', ''))
            else:
                # Create key and model vars for other providers
                key_var_name = f'{provider_id}_key_var'
                if not hasattr(self, key_var_name):
                    setattr(self, key_var_name, tk.StringVar(value=current_config.get(f'{provider_id}_api_key', '')))
                
                model_var_name = f'{provider_id}_model_var'
                if not hasattr(self, model_var_name):
                    default_model = current_config.get(f'{provider_id}_model', self.providers[provider_id]['default_model'])
                    setattr(self, model_var_name, tk.StringVar(value=default_model))
        
        self.create_widgets()
        
    def create_widgets(self):
        """Create settings dialog widgets"""
        # Configure ttk styles for this dialog with current theme
        style = ttk.Style()
        
        style.configure('Settings.TFrame',
                       background=self.colors['bg'])
        
        style.configure('Settings.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'])
        
        style.configure('Settings.Header.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 11, 'bold'))
        
        style.configure('Settings.Title.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 14, 'bold'))
        
        style.configure('Settings.Help.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text_secondary'],
                       font=('Segoe UI', 9))
        
        style.configure('Settings.TEntry',
                       fieldbackground=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'])
        
        style.configure('Settings.TCheckbutton',
                       background=self.colors['bg'],
                       foreground=self.colors['text'])
        
        # Map state-specific colors for checkbuttons (fixes hover in dark mode)
        style.map('Settings.TCheckbutton',
                 background=[('active', self.colors['bg']), ('!active', self.colors['bg'])],
                 foreground=[('active', self.colors['text']), ('!active', self.colors['text'])])
        
        style.configure('Settings.TButton',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'])
        
        style.configure('Settings.TRadiobutton',
                       background=self.colors['bg'],
                       foreground=self.colors['text'])
        
        # Map state-specific colors for radiobuttons (fixes hover in dark mode)
        style.map('Settings.TRadiobutton',
                 background=[('active', self.colors['bg']), ('!active', self.colors['bg'])],
                 foreground=[('active', self.colors['text']), ('!active', self.colors['text'])])
        
        # Create scrollable frame for settings
        canvas = tk.Canvas(self, bg=self.colors['bg'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style='Settings.TFrame')
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        main_frame = ttk.Frame(scrollable_frame, padding="25", style='Settings.TFrame')
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        main_frame.columnconfigure(1, weight=1)
        
        # Title
        title_label = ttk.Label(main_frame,
                               text="Application Settings",
                               style='Settings.Title.TLabel')
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 25))
        
        current_row = 1
        
        # AI Provider Selection
        provider_header = ttk.Label(main_frame,
                                   text="AI Provider Selection",
                                   style='Settings.Header.TLabel')
        provider_header.grid(row=current_row, column=0, columnspan=2, sticky=tk.W, pady=(0, 10))
        current_row += 1
        
        ttk.Label(main_frame, 
                 text="Choose your preferred AI provider:",
                 style='Settings.TLabel').grid(row=current_row, column=0, columnspan=2, sticky=tk.W, pady=5)
        current_row += 1
        
        # Provider radio buttons
        self.provider_var = tk.StringVar(value=self.current_config.get('ai_provider', 'xai'))
        
        for provider_id, provider_info in self.providers.items():
            rb = ttk.Radiobutton(main_frame,
                                text=provider_info['name'],
                                variable=self.provider_var,
                                value=provider_id,
                                style='Settings.TRadiobutton',
                                command=self.on_provider_change)
            rb.grid(row=current_row, column=0, columnspan=2, sticky=tk.W, pady=2, padx=20)
            current_row += 1
        
        current_row += 1
        
        # API Keys Section - Dynamic based on selected provider
        self.api_keys_frame = ttk.Frame(main_frame, style='Settings.TFrame')
        self.api_keys_frame.grid(row=current_row, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(10, 0))
        self.api_keys_frame.columnconfigure(1, weight=1)
        current_row += 1
        
        # Storage for API key entries
        self.api_key_entries = {}
        
        # Create API key fields for all providers
        self.create_api_key_fields()
        
        # Garmin Credentials section
        garmin_header = ttk.Label(main_frame,
                                 text="Garmin Connect Credentials",
                                 style='Settings.Header.TLabel')
        garmin_header.grid(row=current_row, column=0, columnspan=2, sticky=tk.W, pady=(20, 10))
        current_row += 1
        
        ttk.Label(main_frame, text="Email:", style='Settings.TLabel').grid(row=current_row, column=0, sticky=tk.W, pady=8)
        
        self.email_var = tk.StringVar(value=self.current_config.get('garmin_email', ''))
        email_entry = ttk.Entry(main_frame,
                               textvariable=self.email_var,
                               width=50,
                               style='Settings.TEntry')
        email_entry.grid(row=current_row, column=1, sticky=(tk.W, tk.E), pady=8)
        current_row += 1
        
        ttk.Label(main_frame, text="Password:", style='Settings.TLabel').grid(row=current_row, column=0, sticky=tk.W, pady=8)
        
        self.password_var = tk.StringVar(value=self.current_config.get('garmin_password', ''))
        password_entry = ttk.Entry(main_frame,
                                  textvariable=self.password_var,
                                  width=50,
                                  show="*",
                                  style='Settings.TEntry')
        password_entry.grid(row=current_row, column=1, sticky=(tk.W, tk.E), pady=8)
        current_row += 1
        
        # Buttons
        button_frame = ttk.Frame(main_frame, style='Settings.TFrame')
        button_frame.grid(row=current_row, column=0, columnspan=2, pady=(30, 0))
        
        ttk.Button(button_frame,
                  text="Save",
                  command=self.save_settings,
                  style='Settings.TButton').grid(row=0, column=0, padx=5)
        
        ttk.Button(button_frame,
                  text="Cancel",
                  command=self.cancel,
                  style='Settings.TButton').grid(row=0, column=1, padx=5)
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Show initial provider fields
        self.on_provider_change()
    
    def create_api_key_fields(self):
        """Create API key entry fields for all providers"""
        # Clear existing
        for widget in self.api_keys_frame.winfo_children():
            widget.destroy()
        
        row = 0
        
        # Header
        header = ttk.Label(self.api_keys_frame,
                          text="API Configuration",
                          style='Settings.Header.TLabel')
        header.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(0, 10))
        row += 1
        
        selected_provider = self.provider_var.get()
        
        # Show fields based on selected provider
        if selected_provider == 'azure':
            # Azure needs endpoint, deployment, and API key
            ttk.Label(self.api_keys_frame, text="Azure Endpoint:", style='Settings.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
            self.azure_endpoint_var = tk.StringVar(value=self.current_config.get('azure_endpoint', ''))
            ttk.Entry(self.api_keys_frame, textvariable=self.azure_endpoint_var, width=50, style='Settings.TEntry').grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
            row += 1
            
            ttk.Label(self.api_keys_frame, text="Deployment Name:", style='Settings.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
            self.azure_deployment_var = tk.StringVar(value=self.current_config.get('azure_deployment', ''))
            ttk.Entry(self.api_keys_frame, textvariable=self.azure_deployment_var, width=50, style='Settings.TEntry').grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
            row += 1
            
            ttk.Label(self.api_keys_frame, text="API Key:", style='Settings.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
            self.azure_key_var = tk.StringVar(value=self.current_config.get('azure_api_key', ''))
            ttk.Entry(self.api_keys_frame, textvariable=self.azure_key_var, width=50, show="*", style='Settings.TEntry').grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
            row += 1
        else:
            # Standard API key for other providers
            provider_name = self.providers[selected_provider]['name']
            ttk.Label(self.api_keys_frame, text=f"{provider_name} API Key:", style='Settings.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
            
            # Use existing var (already created in __init__)
            key_var = getattr(self, f'{selected_provider}_key_var')
            ttk.Entry(self.api_keys_frame, textvariable=key_var, width=50, show="*", style='Settings.TEntry').grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
            row += 1
            
            # Model selection
            models = self.providers[selected_provider]['models']
            if models:
                ttk.Label(self.api_keys_frame, text="Model:", style='Settings.TLabel').grid(row=row, column=0, sticky=tk.W, pady=5)
                
                # Use existing var (already created in __init__)
                model_var = getattr(self, f'{selected_provider}_model_var')
                model_combo = ttk.Combobox(self.api_keys_frame, textvariable=model_var, values=models, state='readonly', width=47)
                model_combo.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=5)
                row += 1
        
        # Help text
        help_text = self.get_provider_help_text(selected_provider)
        if help_text:
            help_label = ttk.Label(self.api_keys_frame,
                                  text=help_text,
                                  style='Settings.Help.TLabel',
                                  wraplength=550)
            help_label.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(10, 0))
    
    def get_provider_help_text(self, provider):
        """Get help text for each provider"""
        help_texts = {
            'xai': "Get your xAI API key from: https://console.x.ai/",
            'openai': "Get your OpenAI API key from: https://platform.openai.com/api-keys",
            'azure': "Azure endpoint format: https://your-resource.openai.azure.com/",
            'gemini': "Get your Google AI API key from: https://makersuite.google.com/app/apikey",
            'anthropic': "Get your Anthropic API key from: https://console.anthropic.com/"
        }
        return help_texts.get(provider, "")
    
    def on_provider_change(self):
        """Called when provider selection changes"""
        self.create_api_key_fields()
    
    def save_settings(self):
        """Save settings and close dialog"""
        selected_provider = self.provider_var.get()
        
        self.result = {
            'ai_provider': selected_provider,
            'garmin_email': self.email_var.get(),
            'garmin_password': self.password_var.get()
        }
        
        # Save ALL providers' keys (not just selected one)
        # This allows easy switching between providers
        for provider_id in self.providers.keys():
            if provider_id == 'azure':
                # Azure has special fields
                self.result['azure_endpoint'] = self.azure_endpoint_var.get()
                self.result['azure_deployment'] = self.azure_deployment_var.get()
                self.result['azure_api_key'] = self.azure_key_var.get()
            else:
                # Save API key and model for each provider
                key_var = getattr(self, f'{provider_id}_key_var', None)
                if key_var:
                    self.result[f'{provider_id}_api_key'] = key_var.get()
                
                model_var = getattr(self, f'{provider_id}_model_var', None)
                if model_var:
                    self.result[f'{provider_id}_model'] = model_var.get()
        
        self.destroy()
    
    def cancel(self):
        """Cancel and close dialog"""
        self.result = None
        self.destroy()


class GarminChatApp:
    """Main application class for Garmin Chat desktop app"""
    
    def __init__(self, root):
        """Initialize the application"""
        self.root = root
        self.root.title("Garmin Chat")
        self.root.geometry("1200x950")  # Increased from 1000x850 for larger chat area
        
        # Set window icon (works in both script and exe)
        try:
            # Get the correct base path for PyInstaller exe
            if getattr(sys, 'frozen', False):
                # Running as compiled executable
                base_path = Path(sys._MEIPASS)
            else:
                # Running as script
                base_path = Path(__file__).parent
            
            icon_path = base_path / "logo.ico"
            if icon_path.exists():
                self.root.iconbitmap(str(icon_path))
        except Exception as e:
            logger.debug(f"Could not load icon: {e}")
        
        # Set minimum window size
        self.root.minsize(900, 800)  # Increased minimum size too
        
        # Configuration file path
        self.config_dir = Path.home() / ".garmin_chat"
        self.config_dir.mkdir(exist_ok=True)
        self.config_file = self.config_dir / "config.json"
        self.saved_prompts_file = self.config_dir / "saved_prompts.json"
        self.chat_history_dir = self.config_dir / "chat_history"
        self.chat_history_dir.mkdir(exist_ok=True)
        
        # Chat history for current session
        self.current_chat_history = []
        
        # Conversation context memory (last 10 messages for AI context)
        self.conversation_context = []
        self.max_context_messages = 10
        
        # User preferences learned from conversations
        self.user_preferences = {
            'favorite_activities': [],
            'goals': [],
            'interests': [],
            'last_queries': []
        }
        
        # Load conversation history
        self.load_conversation_history()
        
        # Application state
        self.garmin_handler = None
        self.ai_client = None  # Changed from xai_client to ai_client
        self.authenticated = False
        self.mfa_required = False
        
        # AI Provider settings (support multiple providers)
        self.ai_provider = 'xai'  # Default provider
        self.xai_api_key = None
        self.xai_model = 'grok-3'
        self.openai_api_key = None
        self.openai_model = 'gpt-4o'
        self.azure_api_key = None
        self.azure_endpoint = None
        self.azure_deployment = None
        self.gemini_api_key = None
        self.gemini_model = 'gemini-1.5-flash'
        self.anthropic_api_key = None
        self.anthropic_model = 'claude-sonnet-4-5-20250929'
        
        # Garmin settings
        self.garmin_email = None
        self.garmin_password = None
        self.auto_login = True  # Default to auto-login enabled
        self.dark_mode = False  # Start in light mode
        self.window_state_restored = False  # Track if window position was restored
        
        # Load configuration
        self.load_config()
        
        # Configure style
        self.setup_styles()
        
        # Create UI
        self.create_widgets()
        
        # Center window on screen (only if no saved state was restored)
        if not self.window_state_restored:
            self.center_window()
        
        # Set up window close handler to save state
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        # Check if credentials are configured
        current_ai_key = self.get_current_ai_key()
        if not current_ai_key or not self.garmin_email or not self.garmin_password:
            self.root.after(100, self.prompt_for_credentials)
        elif self.auto_login:
            # Auto-connect if credentials are configured and auto-login is enabled
            self.root.after(500, self.auto_connect)
        
    def load_config(self):
        """Load configuration from file"""
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r') as f:
                    config = json.load(f)
                    
                    # AI Provider settings
                    self.ai_provider = config.get('ai_provider', 'xai')
                    self.xai_api_key = config.get('xai_api_key', '')
                    self.xai_model = config.get('xai_model', 'grok-3')
                    self.openai_api_key = config.get('openai_api_key', '')
                    self.openai_model = config.get('openai_model', 'gpt-4o')
                    self.azure_api_key = config.get('azure_api_key', '')
                    self.azure_endpoint = config.get('azure_endpoint', '')
                    self.azure_deployment = config.get('azure_deployment', '')
                    self.gemini_api_key = config.get('gemini_api_key', '')
                    self.gemini_model = config.get('gemini_model', 'gemini-1.5-flash')
                    self.anthropic_api_key = config.get('anthropic_api_key', '')
                    self.anthropic_model = config.get('anthropic_model', 'claude-sonnet-4-5-20250929')
                    
                    # Auto-migrate deprecated model names
                    model_migrations = {
                        # Gemini migrations (experimental models deprecated)
                        'gemini-2.0-flash-exp': 'gemini-1.5-flash',
                        'gemini-exp-1206': 'gemini-1.5-flash',
                        'gemini-1.5-pro-latest': 'gemini-1.5-pro',
                        'gemini-1.5-flash-latest': 'gemini-1.5-flash',
                        # xAI migrations
                        'grok-beta': 'grok-3',
                        'grok-2-1212': 'grok-3',
                        # Add more migrations as models get deprecated
                    }
                    
                    # Migrate Gemini model if deprecated
                    if self.gemini_model in model_migrations:
                        old_model = self.gemini_model
                        self.gemini_model = model_migrations[old_model]
                        logger.info(f"Auto-migrated Gemini model: {old_model} → {self.gemini_model}")
                    
                    # Migrate xAI model if deprecated
                    if self.xai_model in model_migrations:
                        old_model = self.xai_model
                        self.xai_model = model_migrations[old_model]
                        logger.info(f"Auto-migrated xAI model: {old_model} → {self.xai_model}")
                    
                    # Garmin settings
                    self.garmin_email = config.get('garmin_email', '')
                    self.garmin_password = config.get('garmin_password', '')
                    self.auto_login = config.get('auto_login', True)
                    self.dark_mode = config.get('dark_mode', False)
                    
                    # Window state (position and size)
                    window_state = config.get('window_state', {})
                    if window_state:
                        try:
                            width = window_state.get('width', 1200)
                            height = window_state.get('height', 950)
                            x = window_state.get('x')
                            y = window_state.get('y')
                            
                            # Apply saved size
                            if x is not None and y is not None:
                                # Validate position is on screen
                                screen_width = self.root.winfo_screenwidth()
                                screen_height = self.root.winfo_screenheight()
                                
                                # Ensure window is visible
                                if x + width > screen_width:
                                    x = screen_width - width - 10
                                if y + height > screen_height:
                                    y = screen_height - height - 60  # Account for taskbar
                                if x < 0:
                                    x = 10
                                if y < 0:
                                    y = 10
                                
                                self.root.geometry(f'{width}x{height}+{x}+{y}')
                                self.window_state_restored = True
                                logger.info(f"Restored window state: {width}x{height}+{x}+{y}")
                        except Exception as e:
                            logger.warning(f"Could not restore window state: {e}")
                            # Will use center_window() as fallback
                    
                    logger.info("Configuration loaded")
        except Exception as e:
            logger.error(f"Error loading config: {e}")
            self.ai_provider = 'xai'
            self.xai_api_key = None
            self.gemini_model = 'gemini-1.5-flash'
            self.xai_model = 'grok-3'
            self.garmin_email = None
            self.garmin_password = None
            self.auto_login = True
            self.dark_mode = False  # Default to light mode on error
            
    def save_config(self):
        """Save configuration to file"""
        try:
            # Get current window state
            try:
                # Parse geometry string (widthxheight+x+y)
                geometry = self.root.geometry()
                match = geometry.split('+')
                size = match[0].split('x')
                window_state = {
                    'width': int(size[0]),
                    'height': int(size[1]),
                    'x': int(match[1]) if len(match) > 1 else None,
                    'y': int(match[2]) if len(match) > 2 else None
                }
            except:
                window_state = {}
            
            config = {
                # AI Provider settings
                'ai_provider': self.ai_provider,
                'xai_api_key': self.xai_api_key or '',
                'xai_model': self.xai_model,
                'openai_api_key': self.openai_api_key or '',
                'openai_model': self.openai_model,
                'azure_api_key': self.azure_api_key or '',
                'azure_endpoint': self.azure_endpoint or '',
                'azure_deployment': self.azure_deployment or '',
                'gemini_api_key': self.gemini_api_key or '',
                'gemini_model': self.gemini_model,
                'anthropic_api_key': self.anthropic_api_key or '',
                'anthropic_model': self.anthropic_model,
                # Garmin settings
                'garmin_email': self.garmin_email,
                'garmin_password': self.garmin_password,
                'auto_login': self.auto_login,
                'dark_mode': self.dark_mode,
                # Window state
                'window_state': window_state
            }
            with open(self.config_file, 'w') as f:
                json.dump(config, f, indent=2)
            logger.info("Configuration saved")
        except Exception as e:
            logger.error(f"Error saving config: {e}")
    
    def on_closing(self):
        """Handle window close event - save state and exit"""
        try:
            # Save current configuration including window state
            self.save_config()
        except Exception as e:
            logger.error(f"Error saving config on close: {e}")
        finally:
            # Close the application
            self.root.destroy()
            
    def prompt_for_credentials(self):
        """Prompt user to enter credentials on first run"""
        messagebox.showinfo(
            "Welcome to Garmin Chat",
            "Welcome! Before you can start chatting, you need to configure your credentials.\n\n"
            "You'll need:\n"
            "1. An AI API key (xAI, OpenAI, Gemini, Anthropic, or Azure)\n"
            "2. Your Garmin Connect email and password\n\n"
            "Click OK to open the settings dialog.",
            parent=self.root
        )
        self.open_settings()
        
    def auto_connect(self):
        """Automatically connect to Garmin on startup if credentials are configured"""
        self.add_message("System", "Auto-connecting to Garmin Connect...", 'system')
        self.update_status("Connecting to Garmin...", False)
        self.connect_to_garmin()
        
    def setup_styles(self):
        """Configure ttk styles for modern Fluent Design look"""
        # Apply colors based on saved dark_mode preference
        if self.dark_mode:
            # Dark mode colors
            self.colors = {
                'bg': '#202020',
                'card_bg': '#2D2D30',
                'accent': '#60A5FA',
                'accent_hover': '#3B82F6',
                'accent_light': '#1E3A5F',
                'text': '#E5E5E5',
                'text_secondary': '#A0A0A0',
                'border': '#3E3E42',
                'success': '#10B981',
                'warning': '#F59E0B',
                'shadow': '#00000040',
            }
        else:
            # Light mode colors (default)
            self.colors = {
                'bg': '#F3F3F3',            # Light gray background
                'card_bg': '#FFFFFF',        # White cards
                'accent': '#0078D4',         # Windows 11 blue
                'accent_hover': '#106EBE',   # Darker blue on hover
                'accent_light': '#E6F2FA',   # Light blue background
                'text': '#1F1F1F',           # Almost black text
                'text_secondary': '#605E5C', # Gray text
                'border': '#EDEBE9',         # Light border
                'success': '#107C10',        # Green
                'warning': '#D83B01',        # Red/Orange
                'shadow': '#00000010',       # Subtle shadow
            }
        
        # Configure root window
        self.root.configure(bg=self.colors['bg'])
        
        # Modern TTK styles
        style = ttk.Style()
        style.theme_use('clam')
        
        # Base frame style
        style.configure('TFrame',
                       background=self.colors['bg'])
        
        # Card frame style (elevated white cards)
        style.configure('Card.TFrame',
                       background=self.colors['card_bg'],
                       relief='flat',
                       borderwidth=1)
        
        # Modern button styles
        style.configure('Modern.TButton',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat',
                       padding=(12, 6),
                       font=('Segoe UI', 11))  # Larger font for emojis
        style.map('Modern.TButton',
                 background=[('active', self.colors['accent_light'] if self.dark_mode else self.colors['border']), 
                           ('pressed', self.colors['accent_light'] if self.dark_mode else self.colors['border'])],
                 foreground=[('active', self.colors['accent'] if self.dark_mode else self.colors['text']),
                           ('pressed', self.colors['accent'] if self.dark_mode else self.colors['text'])])
        
        # Accent button (primary action)
        style.configure('Accent.TButton',
                       background=self.colors['accent'],
                       foreground='white',
                       borderwidth=0,
                       relief='flat',
                       padding=(16, 8),
                       font=('Segoe UI', 10, 'bold'))
        style.map('Accent.TButton',
                 background=[('active', self.colors['accent_hover']), 
                           ('pressed', self.colors['accent_hover'])],
                 foreground=[('active', 'white'), ('pressed', 'white')])
        
        # Label styles
        style.configure('Title.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 18, 'bold'))
        
        style.configure('Heading.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 11, 'bold'))
        
        style.configure('TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 10))
        
        style.configure('Status.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text_secondary'],
                       font=('Segoe UI', 9))
        
        # Entry style
        style.configure('Modern.TEntry',
                       fieldbackground=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       font=('Segoe UI', 10))
        
        # LabelFrame style (card with title)
        style.configure('Card.TLabelframe',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat')
        style.configure('Card.TLabelframe.Label',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 10, 'bold'))
        
    def center_window(self):
        """Center the window on screen, accounting for taskbar"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        
        # Get screen dimensions
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        
        # Calculate center position
        x = (screen_width // 2) - (width // 2)
        y = (screen_height // 2) - (height // 2)
        
        # Ensure window doesn't go off-screen or under taskbar
        # Account for typical taskbar height (40-50 pixels)
        taskbar_height = 50
        if y + height > screen_height - taskbar_height:
            y = screen_height - height - taskbar_height - 10  # 10px margin
        
        # Ensure window doesn't go above screen
        if y < 0:
            y = 10
        
        # Ensure window doesn't go off sides
        if x < 0:
            x = 10
        if x + width > screen_width:
            x = screen_width - width - 10
        
        self.root.geometry(f'{width}x{height}+{x}+{y}')
        
    def create_widgets(self):
        """Create all UI widgets"""
        
        # Main container with modern background
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(3, weight=1)  # Chat display gets the extra space (now at row 3)
        
        # Row 0: Header Card
        header_card = ttk.Frame(main_frame, style='Card.TFrame', padding="20")
        header_card.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 15))
        header_card.columnconfigure(0, weight=1)
        
        title_label = ttk.Label(header_card, 
                               text="Garmin Chat",
                               style='Title.TLabel')
        title_label.grid(row=0, column=0, sticky=tk.W)
        
        subtitle_label = ttk.Label(header_card,
                                  text="AI-powered insights for your fitness data",
                                  foreground=self.colors['text_secondary'],
                                  background=self.colors['card_bg'],
                                  font=('Segoe UI', 10))
        subtitle_label.grid(row=1, column=0, sticky=tk.W, pady=(5, 0))
        
        # Right side buttons - larger icons
        button_container = ttk.Frame(header_card, style='Card.TFrame')
        button_container.grid(row=0, column=1, rowspan=2, padx=(10, 0))
        
        search_btn = ttk.Button(button_container,
                               text="🔍",
                               command=self.open_search,
                               style='Modern.TButton',
                               width=4)
        search_btn.grid(row=0, column=0, padx=3)
        self.create_tooltip(search_btn, "Search chat history")
        
        theme_btn = ttk.Button(button_container,
                              text="🌙",
                              command=self.toggle_theme,
                              style='Modern.TButton',
                              width=4)
        theme_btn.grid(row=0, column=1, padx=3)
        self.create_tooltip(theme_btn, "Toggle dark mode")
        
        settings_btn = ttk.Button(button_container,
                                 text="Settings",
                                 command=self.open_settings,
                                 style='Modern.TButton')
        settings_btn.grid(row=0, column=2, padx=3)
        self.create_tooltip(settings_btn, "Settings")
        
        # Row 1: Control buttons card
        control_card = ttk.Frame(main_frame, style='Card.TFrame', padding="15")
        control_card.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 15))
        
        self.connect_btn = ttk.Button(control_card,
                                     text="🔐 Connect to Garmin",
                                     command=self.connect_to_garmin,
                                     style='Accent.TButton')
        self.connect_btn.grid(row=0, column=0, padx=(0, 8))
        
        self.refresh_btn = ttk.Button(control_card,
                                     text="Refresh",
                                     command=self.refresh_data,
                                     style='Modern.TButton',
                                     state=tk.DISABLED)
        self.refresh_btn.grid(row=0, column=1, padx=4)
        self.create_tooltip(self.refresh_btn, "Refresh Garmin data")
        
        self.reset_btn = ttk.Button(control_card,
                                   text="🗑️ Reset",
                                   command=self.reset_chat,
                                   style='Modern.TButton',
                                   state=tk.DISABLED)
        self.reset_btn.grid(row=0, column=2, padx=4)
        self.create_tooltip(self.reset_btn, "Clear chat history")
        
        self.save_prompts_btn = ttk.Button(control_card,
                                          text="💾 Prompts",
                                          command=self.open_saved_prompts,
                                          style='Modern.TButton')
        self.save_prompts_btn.grid(row=0, column=3, padx=4)
        self.create_tooltip(self.save_prompts_btn, "Manage saved prompts")
        
        self.save_chat_btn = ttk.Button(control_card,
                                       text="📝 Save",
                                       command=self.save_chat_history,
                                       style='Modern.TButton',
                                       state=tk.DISABLED)
        self.save_chat_btn.grid(row=0, column=4, padx=4)
        self.create_tooltip(self.save_chat_btn, "Save this conversation")
        
        self.view_chats_btn = ttk.Button(control_card,
                                        text="📂 History",
                                        command=self.open_chat_history_viewer,
                                        style='Modern.TButton')
        self.view_chats_btn.grid(row=0, column=5, padx=(4, 0))
        self.create_tooltip(self.view_chats_btn, "View saved chats")
        
        # Favorite button removed - feature was non-functional
        
        self.export_btn = ttk.Button(control_card,
                                     text="📄",
                                     command=self.export_conversation_report,
                                     style='Modern.TButton',
                                     width=4,
                                     state=tk.DISABLED)
        self.export_btn.grid(row=0, column=6, padx=(4, 0))
        self.create_tooltip(self.export_btn, "Export report")
        
        # Status label
        self.status_label = ttk.Label(control_card,
                                     text="●  Not connected",
                                     style='Status.TLabel')
        self.status_label.grid(row=1, column=0, columnspan=7, sticky=tk.W, pady=(10, 0))
        
        # Smart Suggestions and Follow-up Questions removed from main grid
        # They were causing the chat display to shrink when shown
        # These features can be re-enabled later with a better UI approach (popup/sidebar)
        
        # Row 2: MFA card (initially hidden) - moved up from row 4
        # Row 2: MFA card (initially hidden) - moved up from row 4
        self.mfa_frame = ttk.LabelFrame(main_frame, 
                                       text="🔐 Multi-Factor Authentication", 
                                       style='Card.TLabelframe',
                                       padding="20")
        self.mfa_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 15))
        self.mfa_frame.grid_remove()  # Hide initially
        self.mfa_frame.columnconfigure(2, weight=1)
        
        ttk.Label(self.mfa_frame, 
                 text="Enter 6-digit code:", 
                 background=self.colors['card_bg'],
                 foreground=self.colors['text'],
                 font=('Segoe UI', 10)).grid(row=0, column=0, padx=(0, 15), pady=5)
        
        self.mfa_entry = ttk.Entry(self.mfa_frame, 
                                  width=15, 
                                  font=('Segoe UI', 12),
                                  style='Modern.TEntry')
        self.mfa_entry.grid(row=0, column=1, padx=(0, 15), pady=5)
        self.mfa_entry.bind('<Return>', lambda e: self.submit_mfa())
        
        self.mfa_btn = ttk.Button(self.mfa_frame,
                                 text="Submit Code",
                                 command=self.submit_mfa,
                                 style='Accent.TButton')
        self.mfa_btn.grid(row=0, column=2, sticky=tk.W, pady=5)
        
        # Row 3: Chat display card (gets extra space) - moved up from row 5
        chat_card = ttk.Frame(main_frame, style='Card.TFrame', padding="0")
        chat_card.grid(row=3, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 15))
        chat_card.columnconfigure(0, weight=1)
        chat_card.rowconfigure(0, weight=1)
        
        # Chat history (scrolled text) with modern styling
        self.chat_display = scrolledtext.ScrolledText(chat_card,
                                                      wrap=tk.WORD,
                                                      font=('Segoe UI', 11),  # Increased from 10 for better readability
                                                      bg=self.colors['card_bg'],
                                                      fg=self.colors['text'],  # Set text color for dark mode
                                                      relief=tk.FLAT,
                                                      borderwidth=0,
                                                      padx=20,
                                                      pady=15)
        self.chat_display.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.chat_display.config(state=tk.DISABLED)
        
        # Configure text tags with modern colors
        self.chat_display.tag_configure('user', 
                                       foreground=self.colors['accent'], 
                                       font=('Segoe UI', 10, 'bold'))
        self.chat_display.tag_configure('assistant', 
                                       foreground=self.colors['success'], 
                                       font=('Segoe UI', 10, 'bold'))
        self.chat_display.tag_configure('system', 
                                       foreground=self.colors['text_secondary'], 
                                       font=('Segoe UI', 9, 'italic'))
        self.chat_display.tag_configure('timestamp', 
                                       foreground=self.colors['text_secondary'], 
                                       font=('Segoe UI', 8))
        self.chat_display.tag_configure('bold', 
                                       font=('Segoe UI', 10, 'bold'))
        self.chat_display.tag_configure('header', 
                                       font=('Segoe UI', 11, 'bold'), 
                                       foreground=self.colors['text'])
        self.chat_display.tag_configure('table', 
                                       font=('Courier New', 9), 
                                       foreground=self.colors['text'], 
                                       spacing1=2, 
                                       spacing3=2)
        
        # Row 4: Input card (moved up from row 6)
        input_card = ttk.Frame(main_frame, style='Card.TFrame', padding="15")
        input_card.grid(row=4, column=0, sticky=(tk.W, tk.E))
        input_card.columnconfigure(0, weight=1)
        
        # Message input (multi-line Text widget) with modern styling
        self.message_entry = tk.Text(input_card, 
                                     height=3,
                                     font=('Segoe UI', 10),
                                     wrap=tk.WORD,
                                     relief=tk.FLAT,
                                     borderwidth=1,
                                     highlightthickness=1,
                                     highlightbackground=self.colors['border'],
                                     highlightcolor=self.colors['accent'],
                                     bg=self.colors['card_bg'],
                                     fg=self.colors['text'],
                                     insertbackground=self.colors['accent'])
        self.message_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        
        # Send button with modern accent style
        self.send_btn = ttk.Button(input_card,
                                   text="Send →",
                                   command=self.send_message,
                                   state=tk.DISABLED,
                                   style='Accent.TButton')
        self.send_btn.grid(row=0, column=1)
        
        # Helper text
        helper_text = ttk.Label(input_card,
                               text="Ctrl+Enter to send  •  Enter for new line",
                               foreground=self.colors['text_secondary'],
                               background=self.colors['card_bg'],
                               font=('Segoe UI', 8))
        helper_text.grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=(8, 0))
        
        # Bind Ctrl+Enter to send (Enter alone creates new line)
        self.message_entry.bind('<Control-Return>', lambda e: self.send_message())
        self.message_entry.bind('<Control-Key-Return>', lambda e: self.send_message())
        self.message_entry.config(state=tk.DISABLED)
        
        # Row 5: Example questions card (moved up from row 7)
        examples_card = ttk.LabelFrame(main_frame, 
                                      text="Quick Questions", 
                                      style='Card.TLabelframe',
                                      padding="15")
        examples_card.grid(row=5, column=0, sticky=(tk.W, tk.E), pady=(15, 0))
        examples_card.columnconfigure(0, weight=1)
        examples_card.columnconfigure(1, weight=1)
        
        examples = [
            "How many steps did I take today?",
            "What was my last workout?",
            "How did I sleep last night?",
            "Show me my recent activities",
        ]
        
        for i, example in enumerate(examples):
            btn = ttk.Button(examples_card,
                           text=example,
                           style='Modern.TButton',
                           command=lambda q=example: self.use_example(q))
            btn.grid(row=i//2, column=i%2, padx=6, pady=6, sticky=(tk.W, tk.E))
        
    def add_message(self, sender, message, tag='user'):
        """Add a message to the chat display"""
        self.chat_display.config(state=tk.NORMAL)
        
        # Add timestamp
        timestamp = datetime.now().strftime("%H:%M")
        self.chat_display.insert(tk.END, f"[{timestamp}] ", 'timestamp')
        
        # Add sender
        self.chat_display.insert(tk.END, f"{sender}: ", tag)
        
        # Parse and add message with markdown formatting
        if tag == 'assistant':
            self._insert_markdown(message)
        else:
            self.chat_display.insert(tk.END, f"{message}\n\n")
        
        # Save to chat history (but not system messages)
        if tag != 'system':
            self.current_chat_history.append({
                'timestamp': datetime.now().isoformat(),
                'sender': sender,
                'message': message,
                'type': tag
            })
        
        self.chat_display.config(state=tk.DISABLED)
        self.chat_display.see(tk.END)
        
    def _insert_markdown(self, text):
        """Insert text with basic markdown formatting (headers, bold, bullets, tables)"""
        import re
        
        lines = text.split('\n')
        in_table = False
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Detect table (lines with | characters)
            if '|' in line and line.strip().startswith('|'):
                # Skip separator lines (lines with only |, -, and +)
                if re.match(r'^[\|\-\+\s]+$', line):
                    i += 1
                    continue
                
                # This is a table line
                if not in_table:
                    in_table = True
                
                # Render table row with monospace font
                self.chat_display.insert(tk.END, line + '\n', 'table')
                i += 1
                continue
            else:
                # Not a table line
                if in_table:
                    in_table = False
                    self.chat_display.insert(tk.END, '\n')  # Extra space after table
            
            # Handle headers (#### or ### or ## or #)
            if line.startswith('#### '):
                header_text = line[5:]
                self._insert_inline_formatting(header_text)
                self.chat_display.insert(tk.END, '\n', 'header')
            elif line.startswith('### '):
                header_text = line[4:]
                self._insert_inline_formatting(header_text)
                self.chat_display.insert(tk.END, '\n', 'header')
            elif line.startswith('## '):
                header_text = line[3:]
                self._insert_inline_formatting(header_text)
                self.chat_display.insert(tk.END, '\n', 'header')
            elif line.startswith('# '):
                header_text = line[2:]
                self._insert_inline_formatting(header_text)
                self.chat_display.insert(tk.END, '\n', 'header')
            # Handle bullets (- item or * item)
            elif line.strip().startswith(('- ', '* ')):
                bullet_text = '  • ' + line.strip()[2:]
                self._insert_inline_formatting(bullet_text)
                self.chat_display.insert(tk.END, '\n')
            # Handle numbered lists (1. item)
            elif re.match(r'^\d+\.\s', line.strip()):
                self._insert_inline_formatting(line)
                self.chat_display.insert(tk.END, '\n')
            # Regular line with possible inline formatting
            else:
                self._insert_inline_formatting(line)
                self.chat_display.insert(tk.END, '\n')
            
            i += 1
        
        self.chat_display.insert(tk.END, "\n")
    
    def _insert_inline_formatting(self, text):
        """Insert text with inline bold formatting (**text**)"""
        import re
        
        # Split by bold markers **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # Bold text
                bold_text = part[2:-2]
                self.chat_display.insert(tk.END, bold_text, 'bold')
            else:
                # Regular text
                self.chat_display.insert(tk.END, part)
        
    def update_status(self, message, is_error=False):
        """Update the status label"""
        self.status_label.config(text=message)
        if is_error:
            self.status_label.config(foreground='#e74c3c')
        else:
            self.status_label.config(foreground='#27ae60')
            
    def open_settings(self):
        """Open settings dialog"""
        current_config = {
            'ai_provider': self.ai_provider,
            'xai_api_key': self.xai_api_key or '',
            'xai_model': self.xai_model,
            'openai_api_key': self.openai_api_key or '',
            'openai_model': self.openai_model,
            'azure_api_key': self.azure_api_key or '',
            'azure_endpoint': self.azure_endpoint or '',
            'azure_deployment': self.azure_deployment or '',
            'gemini_api_key': self.gemini_api_key or '',
            'gemini_model': self.gemini_model,
            'anthropic_api_key': self.anthropic_api_key or '',
            'anthropic_model': self.anthropic_model,
            'garmin_email': self.garmin_email or '',
            'garmin_password': self.garmin_password or ''
        }
        
        dialog = SettingsDialog(self.root, current_config, colors=self.colors)
        self.root.wait_window(dialog)
        
        if dialog.result:
            # Update AI provider
            self.ai_provider = dialog.result.get('ai_provider', 'xai')
            
            # Update all provider API keys
            self.xai_api_key = dialog.result.get('xai_api_key', '')
            self.xai_model = dialog.result.get('xai_model', 'grok-3')
            self.openai_api_key = dialog.result.get('openai_api_key', '')
            self.openai_model = dialog.result.get('openai_model', 'gpt-4o')
            self.azure_api_key = dialog.result.get('azure_api_key', '')
            self.azure_endpoint = dialog.result.get('azure_endpoint', '')
            self.azure_deployment = dialog.result.get('azure_deployment', '')
            self.gemini_api_key = dialog.result.get('gemini_api_key', '')
            self.gemini_model = dialog.result.get('gemini_model', 'gemini-1.5-flash')
            self.anthropic_api_key = dialog.result.get('anthropic_api_key', '')
            self.anthropic_model = dialog.result.get('anthropic_model', 'claude-sonnet-4-5-20250929')
            
            # Update Garmin credentials
            self.garmin_email = dialog.result.get('garmin_email', '')
            self.garmin_password = dialog.result.get('garmin_password', '')
            
            self.save_config()
            
            # If already authenticated, reinitialize AI client
            if self.ai_client:
                try:
                    self.initialize_ai_client()
                    provider_names = {
                        'xai': 'xAI (Grok)',
                        'openai': 'OpenAI (ChatGPT)',
                        'azure': 'Azure OpenAI',
                        'gemini': 'Google Gemini',
                        'anthropic': 'Anthropic (Claude)'
                    }
                    provider_name = provider_names.get(self.ai_provider, self.ai_provider)
                    self.add_message("System", f"Settings updated! Now using: {provider_name}", 'system')
                except Exception as e:
                    self.add_message("System", f"Error updating AI client: {e}", 'system')
                    
    def initialize_ai_client(self):
        """Initialize AI client based on selected provider"""
        try:
            provider = self.ai_provider
            
            if provider == 'xai' and self.xai_api_key:
                from ai_client import AIClient
                self.ai_client = AIClient(provider='xai', api_key=self.xai_api_key, model=self.xai_model)
                logger.info(f"AI client initialized: xAI ({self.xai_model})")
                return True
                
            elif provider == 'openai' and self.openai_api_key:
                from ai_client import AIClient
                self.ai_client = AIClient(provider='openai', api_key=self.openai_api_key, model=self.openai_model)
                logger.info(f"AI client initialized: OpenAI ({self.openai_model})")
                return True
                
            elif provider == 'azure' and self.azure_api_key and self.azure_endpoint:
                from ai_client import AIClient
                self.ai_client = AIClient(
                    provider='azure',
                    api_key=self.azure_api_key,
                    azure_endpoint=self.azure_endpoint,
                    azure_deployment=self.azure_deployment
                )
                logger.info(f"AI client initialized: Azure OpenAI")
                return True
                
            elif provider == 'gemini' and self.gemini_api_key:
                from ai_client import AIClient
                self.ai_client = AIClient(provider='gemini', api_key=self.gemini_api_key, model=self.gemini_model)
                logger.info(f"AI client initialized: Google Gemini ({self.gemini_model})")
                return True
                
            elif provider == 'anthropic' and self.anthropic_api_key:
                from ai_client import AIClient
                self.ai_client = AIClient(provider='anthropic', api_key=self.anthropic_api_key, model=self.anthropic_model)
                logger.info(f"AI client initialized: Anthropic ({self.anthropic_model})")
                return True
            else:
                logger.warning(f"No valid API key for provider: {provider}")
                return False
                
        except Exception as e:
            logger.error(f"Error initializing AI client: {e}")
            return False
    
    def get_current_ai_key(self):
        """Get the API key for the currently selected provider"""
        provider = self.ai_provider
        if provider == 'xai':
            return self.xai_api_key
        elif provider == 'openai':
            return self.openai_api_key
        elif provider == 'azure':
            return self.azure_api_key
        elif provider == 'gemini':
            return self.gemini_api_key
        elif provider == 'anthropic':
            return self.anthropic_api_key
        return None
    
    def connect_to_garmin(self):
        """Initialize and authenticate with Garmin Connect"""
        # Check if all credentials are configured
        current_ai_key = self.get_current_ai_key()
        
        if not current_ai_key or not self.garmin_email or not self.garmin_password:
            provider_names = {
                'xai': 'xAI',
                'openai': 'OpenAI',
                'azure': 'Azure OpenAI',
                'gemini': 'Google Gemini',
                'anthropic': 'Anthropic'
            }
            provider_name = provider_names.get(self.ai_provider, self.ai_provider)
            
            messagebox.showerror(
                "Configuration Required",
                f"Please configure all your credentials in Settings before connecting to Garmin.\n\n"
                f"You need:\n"
                f"- {provider_name} API key\n"
                f"- Garmin Connect email\n"
                f"- Garmin Connect password",
                parent=self.root
            )
            self.open_settings()
            return
            
        self.connect_btn.config(state=tk.DISABLED)
        self.update_status("Connecting to Garmin...", False)
        
        # Run in thread to prevent UI freezing
        thread = threading.Thread(target=self._authenticate_garmin)
        thread.daemon = True
        thread.start()
        
    def _authenticate_garmin(self):
        """Authenticate with Garmin (runs in thread)"""
        try:
            # Initialize AI client with current provider
            if not self.initialize_ai_client():
                self.root.after(0, lambda: self._on_auth_failure("Failed to initialize AI client. Please check your API key in Settings."))
                return
            
            # Initialize Garmin handler with stored credentials
            self.garmin_handler = GarminDataHandler(self.garmin_email, self.garmin_password)
            result = self.garmin_handler.authenticate()
            
            if result.get('success'):
                self.authenticated = True
                self.mfa_required = False
                self.root.after(0, lambda: self._on_auth_success())
            elif result.get('mfa_required'):
                self.mfa_required = True
                self.authenticated = False
                self.root.after(0, lambda: self._show_mfa_input())
            else:
                error_msg = result.get('error', 'Unknown error')
                self.root.after(0, lambda: self.update_status(f"❌ {error_msg}", True))
                self.root.after(0, lambda: self.connect_btn.config(state=tk.NORMAL))
                
        except Exception as e:
            self.root.after(0, lambda: self.update_status(f"❌ Error: {str(e)}", True))
            self.root.after(0, lambda: self.connect_btn.config(state=tk.NORMAL))
            
    def _show_mfa_input(self):
        """Show MFA input frame"""
        self.mfa_frame.grid()
        self.update_status("🔐 MFA Required: Enter your 6-digit code", False)
        self.mfa_entry.focus()
        
    def submit_mfa(self):
        """Submit MFA code"""
        mfa_code = self.mfa_entry.get().strip()
        
        if not mfa_code or len(mfa_code) != 6:
            self.update_status("❌ Please enter a valid 6-digit MFA code", True)
            return
            
        self.mfa_btn.config(state=tk.DISABLED)
        self.update_status("Submitting MFA code...", False)
        
        # Run in thread
        thread = threading.Thread(target=self._submit_mfa_code, args=(mfa_code,))
        thread.daemon = True
        thread.start()
        
    def _submit_mfa_code(self, mfa_code):
        """Submit MFA code (runs in thread)"""
        try:
            result = self.garmin_handler.submit_mfa(mfa_code)
            
            if result.get('success'):
                self.authenticated = True
                self.mfa_required = False
                self.root.after(0, lambda: self._on_auth_success())
                self.root.after(0, lambda: self.mfa_frame.grid_remove())
            else:
                error_msg = result.get('error', 'Unknown error')
                self.root.after(0, lambda: self.update_status(f"❌ {error_msg}", True))
                self.root.after(0, lambda: self.mfa_btn.config(state=tk.NORMAL))
                
        except Exception as e:
            self.root.after(0, lambda: self.update_status(f"❌ Error: {str(e)}", True))
            self.root.after(0, lambda: self.mfa_btn.config(state=tk.NORMAL))
            
    def _on_auth_success(self):
        """Handle successful authentication"""
        self.update_status("✅ Connected to Garmin Connect!", False)
        self.message_entry.config(state=tk.NORMAL)
        self.send_btn.config(state=tk.NORMAL)
        self.refresh_btn.config(state=tk.NORMAL)
        self.reset_btn.config(state=tk.NORMAL)
        self.save_chat_btn.config(state=tk.NORMAL)
        # favorite_btn removed
        self.export_btn.config(state=tk.NORMAL)
        self.connect_btn.config(state=tk.DISABLED)
        self.message_entry.focus()
        
        self.add_message("System",
                        "Connected! You can now ask questions about your Garmin data.",
                        'system')
        
        # Smart suggestions disabled for better UX - they took up too much vertical space
        # self.root.after(2000, self.show_smart_suggestions)
        
    def send_message(self):
        """Send a message to the chatbot"""
        if not self.authenticated:
            self.update_status("❌ Please connect to Garmin first", True)
            return
            
        # Get message from Text widget
        message = self.message_entry.get("1.0", tk.END).strip()
        if not message:
            return
            
        # Add user message to display
        self.add_message("You", message, 'user')
        
        # Clear input
        self.message_entry.delete("1.0", tk.END)
        
        # Disable input while processing
        self.message_entry.config(state=tk.DISABLED)
        self.send_btn.config(state=tk.DISABLED)
        
        # Process in thread
        thread = threading.Thread(target=self._process_message, args=(message,))
        thread.daemon = True
        thread.start()
        
        # Return 'break' to prevent default behavior when called from key binding
        return 'break'
        
    def _process_message(self, message):
        """Process the message and get AI response (runs in thread)"""
        try:
            # Determine what data to fetch
            query_lower = message.lower()
            
            # Detect if user wants activities by date range or by count
            activity_limit = 5  # Default
            use_date_range = False
            
            # Check for date range requests (last X days/weeks/months)
            import re
            from datetime import datetime, timedelta
            
            # Match "last/past X days/weeks/months" OR "last/this month/week/year"
            time_period_match = re.search(r'(?:last|past)\s+(\d+)\s+(day|week|month)s?', query_lower)
            simple_period_match = re.search(r'(?:last|past|this)\s+(month|week|year)', query_lower)
            
            if time_period_match:
                number = int(time_period_match.group(1))
                unit = time_period_match.group(2)
                
                # Calculate date range
                end_date = datetime.now()
                if unit == "day":
                    start_date = end_date - timedelta(days=number)
                elif unit == "week":
                    start_date = end_date - timedelta(weeks=number)
                elif unit == "month":
                    start_date = end_date - timedelta(days=number * 30)  # Approximate
                
                logger.info(f"Detected date range query: last {number} {unit}(s)")
                logger.info(f"Date range: {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}")
                use_date_range = True
                
            elif simple_period_match:
                # Handle "last month", "this week", etc.
                period = simple_period_match.group(1)
                prefix = simple_period_match.group(0).split()[0]  # "last", "this", or "past"
                
                end_date = datetime.now()
                
                if period == "month":
                    if prefix == "this":
                        # Current month: from 1st to today
                        start_date = end_date.replace(day=1)
                    else:
                        # Last month: 30 days ago
                        start_date = end_date - timedelta(days=30)
                elif period == "week":
                    if prefix == "this":
                        # Current week: last 7 days
                        start_date = end_date - timedelta(days=end_date.weekday())
                    else:
                        # Last week: 7 days ago
                        start_date = end_date - timedelta(days=7)
                elif period == "year":
                    if prefix == "this":
                        # Current year: from Jan 1 to today
                        start_date = end_date.replace(month=1, day=1)
                    else:
                        # Last year: 365 days ago
                        start_date = end_date - timedelta(days=365)
                
                logger.info(f"Detected simple period query: {prefix} {period}")
                logger.info(f"Date range: {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}")
                use_date_range = True
                
            if use_date_range:
                
                # Fetch activities by date range
                try:
                    activities = self.garmin_handler.get_activities_by_date(
                        start_date.strftime("%Y-%m-%d"),
                        end_date.strftime("%Y-%m-%d")
                    )
                    
                    if activities:
                        # Format activities for context
                        context_parts = [f"=== Activities from {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')} ({len(activities)} activities) ==="]
                        for i, activity in enumerate(activities, 1):
                            act_name = activity.get("activityName", "Unknown")
                            act_type = activity.get("activityType", {}).get("typeKey", "Unknown")
                            distance = activity.get("distance", 0) / 1000 if activity.get("distance") else 0
                            duration = activity.get("duration", 0) / 60 if activity.get("duration") else 0
                            calories = activity.get("calories", "N/A")
                            start_time = activity.get("startTimeLocal", "N/A")
                            
                            context_parts.append(f"{i}. {act_name} ({act_type})")
                            context_parts.append(f"   Date: {start_time}")
                            context_parts.append(f"   Distance: {distance:.2f} km")
                            context_parts.append(f"   Duration: {duration:.1f} minutes")
                            context_parts.append(f"   Calories: {calories}")
                            context_parts.append("")
                        
                        garmin_context = "\n".join(context_parts)
                        use_date_range = True
                        logger.info(f"Fetched {len(activities)} activities for date range")
                except Exception as e:
                    logger.error(f"Error fetching activities by date: {e}")
                    # Fall back to regular method
                    use_date_range = False
            
            # If not using date range, detect count-based requests
            if not use_date_range:
                # Check for requests for more activities
                if any(phrase in query_lower for phrase in [
                    "show me more", "more activities", "all activities", "all my activities",
                    "show all", "recent activities"
                ]):
                    activity_limit = 30  # Fetch more for these queries
                    logger.info(f"Detected request for more activities, fetching {activity_limit}")
                
                # Check for specific number requests
                number_match = re.search(r'(?:last|past|recent)\s+(\d+)', query_lower)
                if number_match:
                    requested_count = int(number_match.group(1))
                    activity_limit = min(requested_count, 50)  # Cap at 50
                    logger.info(f"User requested {requested_count} activities, fetching {activity_limit}")
                
                # Fetch appropriate data using regular method
                if any(word in query_lower for word in ["activity", "activities", "workout", "run", "walk", "bike", "exercise"]):
                    garmin_context = self.garmin_handler.format_data_for_context("activities", activity_limit=activity_limit)
                elif any(word in query_lower for word in ["sleep", "rest", "bed"]):
                    garmin_context = self.garmin_handler.format_data_for_context("sleep")
                elif any(word in query_lower for word in ["step", "walk", "distance", "calorie"]):
                    garmin_context = self.garmin_handler.format_data_for_context("summary")
                # NEW: Detect requests for specific health metrics
                elif any(word in query_lower for word in ["body battery", "energy"]):
                    garmin_context = self.garmin_handler.format_data_for_context("body_battery")
                elif any(word in query_lower for word in ["stress", "stressed", "tension"]):
                    garmin_context = self.garmin_handler.format_data_for_context("stress")
                elif any(word in query_lower for word in ["respiration", "breathing", "breath"]):
                    garmin_context = self.garmin_handler.format_data_for_context("respiration")
                elif any(word in query_lower for word in ["hydration", "water", "drink", "fluid"]):
                    garmin_context = self.garmin_handler.format_data_for_context("hydration")
                elif any(word in query_lower for word in ["nutrition", "food", "eat", "meal", "diet", "protein", "carbs", "fat", "macros", "calories consumed", "food log", "logged"]):
                    garmin_context = self.garmin_handler.format_data_for_context("nutrition")
                elif any(word in query_lower for word in ["floor", "climb", "stairs", "elevation"]):
                    garmin_context = self.garmin_handler.format_data_for_context("floors")
                elif any(word in query_lower for word in ["intensity", "vigorous", "moderate"]):
                    garmin_context = self.garmin_handler.format_data_for_context("intensity")
                elif any(word in query_lower for word in ["spo2", "oxygen", "pulse ox"]):
                    garmin_context = self.garmin_handler.format_data_for_context("spo2")
                elif any(word in query_lower for word in ["hrv", "heart rate variability", "variability"]):
                    garmin_context = self.garmin_handler.format_data_for_context("hrv")
                elif any(word in query_lower for word in ["vo2", "fitness age", "training status", "training load"]):
                    garmin_context = self.garmin_handler.format_data_for_context("training")
                # Comprehensive health overview
                elif any(word in query_lower for word in ["health", "wellness", "overview", "summary"]):
                    garmin_context = self.garmin_handler.format_data_for_context("comprehensive")
                else:
                    garmin_context = self.garmin_handler.format_data_for_context("all", activity_limit=activity_limit)
            
            # Add conversation context for memory
            context_summary = ""
            if self.conversation_context:
                recent_convs = self.conversation_context[-5:]
                context_summary = "\n\nPrevious conversation context:\n"
                for conv in recent_convs:
                    sender = conv.get('sender', 'User')
                    msg = conv.get('message', '')[:100]  # First 100 chars
                    context_summary += f"{sender}: {msg}...\n"
            
            enhanced_context = garmin_context + context_summary
            
            # Get AI response
            response = self.ai_client.chat(message, enhanced_context)
            
            # Add response to display
            self.root.after(0, lambda: self.add_message("Garmin Chat", response, 'assistant'))
            
            # Follow-up suggestions disabled for better UX - they took up too much vertical space
            # self.root.after(0, lambda: self.show_followup_buttons(response))
            
            # Update conversation context
            self.conversation_context.append({
                'sender': 'You',
                'message': message,
                'timestamp': datetime.now().isoformat()
            })
            self.conversation_context.append({
                'sender': 'Garmin Chat',
                'message': response,
                'timestamp': datetime.now().isoformat()
            })
            
            # Keep context manageable
            if len(self.conversation_context) > self.max_context_messages:
                self.conversation_context = self.conversation_context[-self.max_context_messages:]
            
        except Exception as e:
            error_msg = f"Sorry, I encountered an error: {str(e)}"
            self.root.after(0, lambda: self.add_message("System", error_msg, 'system'))
            
        finally:
            # Re-enable input
            self.root.after(0, lambda: self.message_entry.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.send_btn.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.message_entry.focus())
            
    def use_example(self, question):
        """Use an example question"""
        if not self.authenticated:
            self.update_status("❌ Please connect to Garmin first", True)
            return
            
        self.message_entry.delete("1.0", tk.END)
        self.message_entry.insert("1.0", question)
        self.send_message()
        
    def refresh_data(self):
        """Refresh Garmin data"""
        self.refresh_btn.config(state=tk.DISABLED)
        self.update_status("Refreshing data...", False)
        
        thread = threading.Thread(target=self._refresh_data)
        thread.daemon = True
        thread.start()
        
    def _refresh_data(self):
        """Refresh data (runs in thread)"""
        try:
            result = self.garmin_handler.authenticate()
            if result.get('success'):
                self.root.after(0, lambda: self.update_status("✅ Data refreshed!", False))
                self.root.after(0, lambda: self.add_message("System", "Data refreshed successfully!", 'system'))
            elif result.get('mfa_required'):
                # MFA is required for refresh
                self.mfa_required = True
                self.authenticated = False
                self.root.after(0, lambda: self._show_mfa_input())
                self.root.after(0, lambda: self.update_status("🔐 MFA Required: Enter your 6-digit code", False))
            else:
                error_msg = result.get('error', 'Unknown error')
                self.root.after(0, lambda: self.update_status(f"❌ {error_msg}", True))
        except Exception as e:
            self.root.after(0, lambda: self.update_status(f"❌ Error: {str(e)}", True))
        finally:
            self.root.after(0, lambda: self.refresh_btn.config(state=tk.NORMAL))
            
    def reset_chat(self):
        """Reset the conversation"""
        if self.ai_client:
            self.ai_client.reset_conversation()
            
        self.chat_display.config(state=tk.NORMAL)
        self.chat_display.delete(1.0, tk.END)
        self.chat_display.config(state=tk.DISABLED)
        self.current_chat_history = []
        
        self.add_message("System", "Conversation reset!", 'system')
        self.update_status("✅ Chat reset", False)
    
    def open_saved_prompts(self):
        """Open dialog to manage saved prompts"""
        SavedPromptsDialog(self.root, self)
    
    def load_saved_prompts(self):
        """Load saved prompts from file"""
        try:
            if self.saved_prompts_file.exists():
                with open(self.saved_prompts_file, 'r') as f:
                    return json.load(f)
            return []
        except Exception as e:
            logger.error(f"Error loading saved prompts: {e}")
            return []
    
    def save_prompt(self, name, prompt):
        """Save a prompt for reuse"""
        try:
            prompts = self.load_saved_prompts()
            prompts.append({'name': name, 'prompt': prompt, 'created': datetime.now().isoformat()})
            with open(self.saved_prompts_file, 'w') as f:
                json.dump(prompts, f, indent=2)
            logger.info(f"Saved prompt: {name}")
        except Exception as e:
            logger.error(f"Error saving prompt: {e}")
    
    def delete_saved_prompt(self, index):
        """Delete a saved prompt"""
        try:
            prompts = self.load_saved_prompts()
            if 0 <= index < len(prompts):
                deleted = prompts.pop(index)
                with open(self.saved_prompts_file, 'w') as f:
                    json.dump(prompts, f, indent=2)
                logger.info(f"Deleted prompt: {deleted['name']}")
        except Exception as e:
            logger.error(f"Error deleting prompt: {e}")
    
    def save_chat_history(self):
        """Save current chat session to file"""
        if not self.current_chat_history:
            messagebox.showinfo("No Chat History", "There's no chat history to save yet!", parent=self.root)
            return
        
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = self.chat_history_dir / f"chat_{timestamp}.json"
            
            with open(filename, 'w') as f:
                json.dump({
                    'saved_at': datetime.now().isoformat(),
                    'messages': self.current_chat_history
                }, f, indent=2)
            
            messagebox.showinfo("Chat Saved", f"Chat history saved successfully!\n\nLocation: {filename}", parent=self.root)
            logger.info(f"Saved chat history to: {filename}")
        except Exception as e:
            logger.error(f"Error saving chat history: {e}")
            messagebox.showerror("Save Error", f"Failed to save chat history: {e}", parent=self.root)
    
    def open_chat_history_viewer(self):
        """Open dialog to view saved chat histories"""
        ChatHistoryViewer(self.root, self)
    
    def load_conversation_history(self):
        """Load recent conversation context for AI memory"""
        try:
            # Load last 5 chat files for context
            chat_files = sorted(self.chat_history_dir.glob("chat_*.json"),
                              key=lambda f: f.stat().st_mtime,
                              reverse=True)[:5]
            
            for file in chat_files:
                with open(file, 'r') as f:
                    data = json.load(f)
                    messages = data.get('messages', [])
                    # Add to context (last message from each chat)
                    if messages:
                        self.conversation_context.extend(messages[-3:])  # Last 3 from each
            
            # Keep only most recent messages
            self.conversation_context = self.conversation_context[-self.max_context_messages:]
        except Exception as e:
            logger.error(f"Error loading conversation history: {e}")
    
    def toggle_theme(self):
        """Toggle between light and dark mode"""
        # Check current theme
        if not hasattr(self, 'dark_mode'):
            self.dark_mode = False
        
        self.dark_mode = not self.dark_mode
        
        if self.dark_mode:
            # Dark mode colors
            self.colors = {
                'bg': '#202020',
                'card_bg': '#2D2D30',
                'accent': '#60A5FA',
                'accent_hover': '#3B82F6',
                'accent_light': '#1E3A5F',
                'text': '#E5E5E5',
                'text_secondary': '#A0A0A0',
                'border': '#3E3E42',
                'success': '#10B981',
                'warning': '#F59E0B',
                'shadow': '#00000040',
            }
        else:
            # Light mode colors (original)
            self.colors = {
                'bg': '#F3F3F3',
                'card_bg': '#FFFFFF',
                'accent': '#0078D4',
                'accent_hover': '#106EBE',
                'accent_light': '#E6F2FA',
                'text': '#1F1F1F',
                'text_secondary': '#605E5C',
                'border': '#EDEBE9',
                'success': '#107C10',
                'warning': '#D83B01',
                'shadow': '#00000010',
            }
        
        # Apply new theme immediately
        self.apply_theme()
        
        # Save theme preference
        self.save_config()
        logger.info(f"Theme toggled to {'dark' if self.dark_mode else 'light'} mode and saved")
    
    def apply_theme(self):
        """Apply current theme colors to all UI elements"""
        # Update root background
        self.root.configure(bg=self.colors['bg'])
        
        # Re-configure ttk styles with new colors
        style = ttk.Style()
        
        # Base frame style
        style.configure('TFrame',
                       background=self.colors['bg'])
        
        # Card frame style
        style.configure('Card.TFrame',
                       background=self.colors['card_bg'],
                       relief='flat',
                       borderwidth=1)
        
        # Label styles
        style.configure('Title.TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 18, 'bold'))
        
        style.configure('Heading.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 11, 'bold'))
        
        style.configure('TLabel',
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 10))
        
        style.configure('Status.TLabel',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text_secondary'],
                       font=('Segoe UI', 9))
        
        # Button styles
        style.configure('Modern.TButton',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat',
                       padding=(12, 6),
                       font=('Segoe UI', 11))
        style.map('Modern.TButton',
                 background=[('active', self.colors['accent_light'] if self.dark_mode else self.colors['border']), 
                           ('pressed', self.colors['accent_light'] if self.dark_mode else self.colors['border'])],
                 foreground=[('active', self.colors['accent'] if self.dark_mode else self.colors['text']),
                           ('pressed', self.colors['accent'] if self.dark_mode else self.colors['text'])])
        
        style.configure('Accent.TButton',
                       background=self.colors['accent'],
                       foreground='white',
                       borderwidth=0,
                       relief='flat',
                       padding=(16, 8),
                       font=('Segoe UI', 10, 'bold'))
        style.map('Accent.TButton',
                 background=[('active', self.colors['accent_hover']), 
                           ('pressed', self.colors['accent_hover'])],
                 foreground=[('active', 'white'), ('pressed', 'white')])
        
        # LabelFrame style
        style.configure('Card.TLabelframe',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       bordercolor=self.colors['border'],
                       borderwidth=1,
                       relief='flat')
        style.configure('Card.TLabelframe.Label',
                       background=self.colors['card_bg'],
                       foreground=self.colors['text'],
                       font=('Segoe UI', 10, 'bold'))
        
        # Update chat display
        self.chat_display.config(bg=self.colors['card_bg'], fg=self.colors['text'])
        
        # Update chat display tags
        self.chat_display.tag_configure('user', foreground=self.colors['accent'])
        self.chat_display.tag_configure('assistant', foreground=self.colors['success'])
        self.chat_display.tag_configure('system', foreground=self.colors['text_secondary'])
        self.chat_display.tag_configure('timestamp', foreground=self.colors['text_secondary'])
        self.chat_display.tag_configure('header', foreground=self.colors['text'])
        self.chat_display.tag_configure('table', foreground=self.colors['text'])
        
        # Update message entry
        self.message_entry.config(
            bg=self.colors['card_bg'],
            fg=self.colors['text'],
            insertbackground=self.colors['accent'],
            highlightbackground=self.colors['border'],
            highlightcolor=self.colors['accent']
        )
        
        # Update MFA entry if it exists
        try:
            self.mfa_entry.config(
                fieldbackground=self.colors['card_bg'],
                foreground=self.colors['text']
            )
        except:
            pass
        
        # Update suggestions label
        try:
            self.suggestions_label.config(
                background=self.colors['card_bg'],
                foreground=self.colors['text_secondary']
            )
        except:
            pass
        
        # Update all frames and labels recursively
        self._update_widget_colors(self.root)
    
    def _update_widget_colors(self, widget):
        """Recursively update widget colors"""
        try:
            widget_type = widget.winfo_class()
            
            # Update Frame backgrounds
            if widget_type in ('TFrame', 'Frame'):
                try:
                    widget.configure(background=self.colors['bg'])
                except:
                    pass
            
            # Update Label colors
            elif widget_type in ('TLabel', 'Label'):
                try:
                    widget.configure(
                        background=self.colors['bg'],
                        foreground=self.colors['text']
                    )
                except:
                    pass
            
            # Update specific styled labels
            if hasattr(widget, 'cget'):
                try:
                    # Check if it's a card background element
                    if 'Card' in str(widget.winfo_parent()):
                        widget.configure(background=self.colors['card_bg'])
                except:
                    pass
            
            # Recursively update children
            for child in widget.winfo_children():
                self._update_widget_colors(child)
        
        except:
            pass
    
    def open_search(self):
        """Open search dialog for chat history"""
        SearchDialog(self.root, self)
    
    # toggle_favorite_chat method removed - feature was non-functional
    
    def show_smart_suggestions(self):
        """Generate and display smart suggestions based on user data"""
        if not self.authenticated:
            return
        
        # Show suggestions frame
        self.suggestions_frame.grid()
        
        # Generate suggestions (this would analyze actual data)
        suggestions = []
        
        # Check when they last asked about certain topics
        recent_topics = [msg.get('message', '').lower() for msg in self.current_chat_history[-10:]]
        
        if not any('sleep' in topic for topic in recent_topics):
            suggestions.append("You haven't checked your sleep data recently")
        
        if not any('steps' in topic or 'walking' in topic for topic in recent_topics):
            suggestions.append("How about reviewing your step count?")
        
        if not any('heart' in topic for topic in recent_topics):
            suggestions.append("Check your heart rate trends")
        
        if suggestions:
            suggestion_text = "💡 " + " • ".join(suggestions[:2])
            self.suggestions_label.config(text=suggestion_text)
        else:
            self.suggestions_frame.grid_remove()
    
    def show_followup_buttons(self, response_text):
        """Show context-aware follow-up buttons after AI response"""
        # Clear existing buttons
        for widget in self.followup_frame.winfo_children():
            widget.destroy()
        
        # Generate follow-up questions based on response
        followups = []
        
        response_lower = response_text.lower()
        
        if 'steps' in response_lower or 'walking' in response_lower:
            followups = [
                "Compare to last week",
                "Show me a weekly trend",
                "What's my daily average?"
            ]
        elif 'sleep' in response_lower:
            followups = [
                "How does this compare to my goal?",
                "Show sleep quality trends",
                "What affects my sleep?"
            ]
        elif 'workout' in response_lower or 'activity' in response_lower:
            followups = [
                "Show workout details",
                "Compare to previous workouts",
                "What's my weekly total?"
            ]
        elif 'heart' in response_lower:
            followups = [
                "Show resting heart rate trend",
                "Compare to healthy range",
                "What's my max heart rate?"
            ]
        else:
            followups = [
                "Tell me more",
                "Show details",
                "Any recommendations?"
            ]
        
        if followups:
            self.followup_frame.grid()
            
            ttk.Label(self.followup_frame,
                     text="Quick follow-ups:",
                     background=self.colors['card_bg'],
                     foreground=self.colors['text_secondary'],
                     font=('Segoe UI', 9)).grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
            
            for i, question in enumerate(followups[:3]):
                btn = ttk.Button(self.followup_frame,
                               text=question,
                               style='Modern.TButton',
                               command=lambda q=question: self.use_example(q))
                btn.grid(row=0, column=i+1, padx=5)
        else:
            self.followup_frame.grid_remove()
    
    def export_conversation_report(self):
        """Export current conversation as a formatted document"""
        ExportReportDialog(self.root, self)
    
    def create_tooltip(self, widget, text):
        """Create a tooltip for a widget"""
        tooltip = ToolTip(widget, text)


class ToolTip:
    """Simple tooltip class"""
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip_window = None
        self.widget.bind('<Enter>', self.show_tooltip)
        self.widget.bind('<Leave>', self.hide_tooltip)
    
    def show_tooltip(self, event=None):
        """Display tooltip"""
        if self.tooltip_window or not self.text:
            return
        
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5
        
        self.tooltip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        
        label = tk.Label(tw, text=self.text, 
                        background="#2D2D30",
                        foreground="#E5E5E5",
                        relief=tk.SOLID,
                        borderwidth=1,
                        font=('Segoe UI', 9),
                        padx=8,
                        pady=4)
        label.pack()
    
    def hide_tooltip(self, event=None):
        """Hide tooltip"""
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None


class SavedPromptsDialog(tk.Toplevel):
    """Dialog for managing saved prompts"""
    
    def __init__(self, parent, app):
        super().__init__(parent)
        self.title("Saved Prompts")
        self.geometry("700x500")
        self.app = app
        
        # Make modal
        self.transient(parent)
        self.grab_set()
        
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)
        
        # Title
        title = ttk.Label(main_frame, text="Saved Prompts", font=('Segoe UI', 14, 'bold'))
        title.grid(row=0, column=0, sticky=tk.W, pady=(0, 15))
        
        # Prompts list
        list_frame = ttk.Frame(main_frame)
        list_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 15))
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(0, weight=1)
        
        # Scrolled listbox
        self.prompts_listbox = tk.Listbox(list_frame, font=('Segoe UI', 10), height=15)
        self.prompts_listbox.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.prompts_listbox.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.prompts_listbox.config(yscrollcommand=scrollbar.set)
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=2, column=0, pady=(0, 10))
        
        ttk.Button(button_frame, text="➕ New Prompt", command=self.new_prompt).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="✏️ Use Selected", command=self.use_prompt).grid(row=0, column=1, padx=5)
        ttk.Button(button_frame, text="🗑️ Delete", command=self.delete_prompt).grid(row=0, column=2, padx=5)
        ttk.Button(button_frame, text="Close", command=self.destroy).grid(row=0, column=3, padx=5)
        
        self.load_prompts()
    
    def load_prompts(self):
        """Load and display saved prompts"""
        self.prompts_listbox.delete(0, tk.END)
        self.prompts = self.app.load_saved_prompts()
        
        for prompt in self.prompts:
            display = f"{prompt['name']} - {prompt['prompt'][:50]}..."
            self.prompts_listbox.insert(tk.END, display)
    
    def new_prompt(self):
        """Create new saved prompt"""
        dialog = tk.Toplevel(self)
        dialog.title("New Prompt")
        dialog.geometry("500x300")
        dialog.transient(self)
        dialog.grab_set()
        
        frame = ttk.Frame(dialog, padding="20")
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        dialog.columnconfigure(0, weight=1)
        dialog.rowconfigure(0, weight=1)
        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(2, weight=1)
        
        ttk.Label(frame, text="Prompt Name:", font=('Segoe UI', 10)).grid(row=0, column=0, sticky=tk.W, pady=5)
        name_entry = ttk.Entry(frame, font=('Segoe UI', 10))
        name_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=5, padx=(10, 0))
        
        ttk.Label(frame, text="Prompt Text:", font=('Segoe UI', 10)).grid(row=1, column=0, sticky=(tk.W, tk.N), pady=5)
        prompt_text = tk.Text(frame, font=('Segoe UI', 10), wrap=tk.WORD, height=8)
        prompt_text.grid(row=1, column=1, rowspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5, padx=(10, 0))
        
        def save():
            name = name_entry.get().strip()
            prompt = prompt_text.get("1.0", tk.END).strip()
            
            if not name or not prompt:
                messagebox.showerror("Validation Error", "Please enter both name and prompt text", parent=dialog)
                return
            
            self.app.save_prompt(name, prompt)
            self.load_prompts()
            dialog.destroy()
        
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=(15, 0))
        
        ttk.Button(button_frame, text="Save", command=save).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Cancel", command=dialog.destroy).grid(row=0, column=1, padx=5)
    
    def use_prompt(self):
        """Use selected prompt"""
        selection = self.prompts_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a prompt to use", parent=self)
            return
        
        index = selection[0]
        prompt = self.prompts[index]
        
        # Insert into message entry and close dialog
        self.app.message_entry.delete("1.0", tk.END)
        self.app.message_entry.insert("1.0", prompt['prompt'])
        self.destroy()
    
    def delete_prompt(self):
        """Delete selected prompt"""
        selection = self.prompts_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a prompt to delete", parent=self)
            return
        
        index = selection[0]
        prompt = self.prompts[index]
        
        if messagebox.askyesno("Confirm Delete", f"Delete prompt '{prompt['name']}'?", parent=self):
            self.app.delete_saved_prompt(index)
            self.load_prompts()






class SearchDialog(tk.Toplevel):
    """Dialog for searching chat history"""
    
    def __init__(self, parent, app):
        super().__init__(parent)
        self.title("Search Chat History")
        self.geometry("700x500")
        self.app = app
        
        self.transient(parent)
        self.grab_set()
        
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(2, weight=1)
        
        # Title
        ttk.Label(main_frame, text="🔍 Search Chat History", 
                 font=('Segoe UI', 14, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=(0, 15))
        
        # Search box
        search_frame = ttk.Frame(main_frame)
        search_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 15))
        search_frame.columnconfigure(0, weight=1)
        
        self.search_entry = ttk.Entry(search_frame, font=('Segoe UI', 10))
        self.search_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        self.search_entry.bind('<Return>', lambda e: self.perform_search())
        
        ttk.Button(search_frame, text="Search", 
                  command=self.perform_search).grid(row=0, column=1)
        
        # Results
        results_frame = ttk.Frame(main_frame)
        results_frame.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(0, weight=1)
        
        self.results_text = scrolledtext.ScrolledText(results_frame,
                                                      wrap=tk.WORD,
                                                      font=('Segoe UI', 9),
                                                      state=tk.DISABLED)
        self.results_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Close button
        ttk.Button(main_frame, text="Close", 
                  command=self.destroy).grid(row=3, column=0, pady=(15, 0))
    
    def perform_search(self):
        """Search through all saved chats"""
        query = self.search_entry.get().strip().lower()
        if not query:
            return
        
        self.results_text.config(state=tk.NORMAL)
        self.results_text.delete(1.0, tk.END)
        
        try:
            chat_files = sorted(self.app.chat_history_dir.glob("chat_*.json"),
                              key=lambda f: f.stat().st_mtime,
                              reverse=True)
            
            results_found = 0
            for file in chat_files:
                with open(file, 'r') as f:
                    data = json.load(f)
                    messages = data.get('messages', [])
                    
                    for msg in messages:
                        text = msg.get('message', '').lower()
                        if query in text:
                            results_found += 1
                            sender = msg.get('sender', 'Unknown')
                            timestamp = msg.get('timestamp', '')
                            
                            self.results_text.insert(tk.END, f"\n{sender} ({timestamp[:10]}):\n", 'bold')
                            self.results_text.insert(tk.END, f"{msg.get('message', '')}...\n\n")
                            
                            if results_found >= 20:  # Limit results
                                break
                
                if results_found >= 20:
                    break
            
            if results_found == 0:
                self.results_text.insert(tk.END, f"No results found for '{query}'")
            else:
                self.results_text.insert(tk.END, f"\n--- {results_found} results found ---")
        
        except Exception as e:
            self.results_text.insert(tk.END, f"Error searching: {e}")
        
        self.results_text.config(state=tk.DISABLED)


class ExportReportDialog(tk.Toplevel):
    """Dialog for exporting conversation reports"""
    
    def __init__(self, parent, app):
        super().__init__(parent)
        self.title("Export Report")
        self.geometry("500x400")
        self.app = app
        
        self.transient(parent)
        self.grab_set()
        
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        
        # Title
        ttk.Label(main_frame, text="📄 Export Report", 
                 font=('Segoe UI', 14, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=(0, 20))
        
        # Export options
        ttk.Label(main_frame, text="Choose export format:", 
                 font=('Segoe UI', 10)).grid(row=1, column=0, sticky=tk.W, pady=(0, 10))
        
        self.export_format = tk.StringVar(value="pdf")
        
        ttk.Radiobutton(main_frame, text="PDF Document", 
                       variable=self.export_format, value="pdf").grid(row=2, column=0, sticky=tk.W, pady=3)
        ttk.Radiobutton(main_frame, text="Word Document (.docx)", 
                       variable=self.export_format, value="docx").grid(row=3, column=0, sticky=tk.W, pady=3)
        ttk.Radiobutton(main_frame, text="Text File (.txt)", 
                       variable=self.export_format, value="txt").grid(row=4, column=0, sticky=tk.W, pady=3)
        
        # Include options
        ttk.Label(main_frame, text="\nInclude:", 
                 font=('Segoe UI', 10, 'bold')).grid(row=5, column=0, sticky=tk.W, pady=(15, 10))
        
        self.include_timestamp = tk.BooleanVar(value=True)
        self.include_system = tk.BooleanVar(value=False)
        
        ttk.Checkbutton(main_frame, text="Timestamps", 
                       variable=self.include_timestamp).grid(row=6, column=0, sticky=tk.W, pady=3)
        ttk.Checkbutton(main_frame, text="System messages", 
                       variable=self.include_system).grid(row=7, column=0, sticky=tk.W, pady=3)
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=8, column=0, pady=(20, 0))
        
        ttk.Button(button_frame, text="Export", 
                  command=self.export_report).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Cancel", 
                  command=self.destroy).grid(row=0, column=1, padx=5)
    
    def export_report(self):
        """Export the conversation in selected format"""
        if not self.app.current_chat_history:
            messagebox.showwarning("No Data", "No conversation to export!", parent=self)
            return
        
        format_type = self.export_format.get()
        
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            if format_type == "txt":
                filename = self.app.chat_history_dir / f"report_{timestamp}.txt"
                self._export_txt(filename)
            elif format_type == "pdf":
                filename = self.app.chat_history_dir / f"report_{timestamp}.pdf"
                self._export_pdf(filename)
            elif format_type == "docx":
                filename = self.app.chat_history_dir / f"report_{timestamp}.docx"
                self._export_docx(filename)
            
            messagebox.showinfo("Export Complete", 
                              f"Report exported successfully!\n\n{filename}",
                              parent=self)
            self.destroy()
        
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export: {e}", parent=self)
    
    def _export_txt(self, filename):
        """Export as plain text"""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("GARMIN CHAT CONVERSATION REPORT\n")
            f.write("=" * 60 + "\n\n")
            
            for msg in self.app.current_chat_history:
                if not self.include_system.get() and msg.get('type') == 'system':
                    continue
                
                sender = msg.get('sender', 'Unknown')
                text = msg.get('message', '')
                
                if self.include_timestamp.get():
                    timestamp = msg.get('timestamp', '')
                    f.write(f"[{timestamp}] {sender}:\n")
                else:
                    f.write(f"{sender}:\n")
                
                f.write(f"{text}\n\n")
                f.write("-" * 60 + "\n\n")
    
    def _export_pdf(self, filename):
        """Export as PDF"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            
            doc = SimpleDocTemplate(str(filename), pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title = Paragraph("Garmin Chat Conversation Report", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 0.3*inch))
            
            # Messages
            for msg in self.app.current_chat_history:
                if not self.include_system.get() and msg.get('type') == 'system':
                    continue
                
                sender = msg.get('sender', 'Unknown')
                text = msg.get('message', '')
                
                if self.include_timestamp.get():
                    timestamp = msg.get('timestamp', '')
                    header = f"<b>[{timestamp}] {sender}:</b>"
                else:
                    header = f"<b>{sender}:</b>"
                
                story.append(Paragraph(header, styles['Normal']))
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            doc.build(story)
        
        except ImportError:
            # Fallback to text if reportlab not installed
            messagebox.showwarning("PDF Export", 
                                 "PDF export requires 'reportlab' package.\nExporting as text instead.",
                                 parent=self)
            txt_filename = str(filename).replace('.pdf', '.txt')
            self._export_txt(Path(txt_filename))
    
    def _export_docx(self, filename):
        """Export as Word document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Garmin Chat Conversation Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Messages
            for msg in self.app.current_chat_history:
                if not self.include_system.get() and msg.get('type') == 'system':
                    continue
                
                sender = msg.get('sender', 'Unknown')
                text = msg.get('message', '')
                
                if self.include_timestamp.get():
                    timestamp = msg.get('timestamp', '')
                    p = doc.add_paragraph()
                    p.add_run(f"[{timestamp}] {sender}:").bold = True
                else:
                    p = doc.add_paragraph()
                    p.add_run(f"{sender}:").bold = True
                
                doc.add_paragraph(text)
                doc.add_paragraph("_" * 60)
            
            doc.save(str(filename))
        
        except ImportError:
            # Fallback to text if python-docx not installed
            messagebox.showwarning("Word Export", 
                                 "Word export requires 'python-docx' package.\nExporting as text instead.",
                                 parent=self)
            txt_filename = str(filename).replace('.docx', '.txt')
            self._export_txt(Path(txt_filename))


class ChatHistoryViewer(tk.Toplevel):
    """Dialog for viewing saved chat histories"""
    
    def __init__(self, parent, app):
        super().__init__(parent)
        self.title("Chat History")
        self.geometry("900x650")
        self.app = app
        
        # Make modal
        self.transient(parent)
        self.grab_set()
        
        # Main container with two columns
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(1, weight=1)
        
        # Title
        title = ttk.Label(main_frame, text="Saved Chat History", font=('Segoe UI', 14, 'bold'))
        title.grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 15))
        
        # Left panel - Chat list
        left_frame = ttk.LabelFrame(main_frame, text="Saved Chats", padding="10")
        left_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 10))
        left_frame.columnconfigure(0, weight=1)
        left_frame.rowconfigure(0, weight=1)
        
        # Listbox with scrollbar
        list_frame = ttk.Frame(left_frame)
        list_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(0, weight=1)
        
        self.chat_listbox = tk.Listbox(list_frame, font=('Segoe UI', 9), width=30)
        self.chat_listbox.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.chat_listbox.bind('<<ListboxSelect>>', self.on_chat_select)
        
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.chat_listbox.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.chat_listbox.config(yscrollcommand=scrollbar.set)
        
        # Right panel - Chat viewer
        right_frame = ttk.LabelFrame(main_frame, text="Chat Content", padding="10")
        right_frame.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S))
        right_frame.columnconfigure(0, weight=1)
        right_frame.rowconfigure(1, weight=1)
        
        # Chat info
        self.info_label = ttk.Label(right_frame, text="Select a chat to view", font=('Segoe UI', 9))
        self.info_label.grid(row=0, column=0, sticky=tk.W, pady=(0, 5))
        
        # Chat content display
        self.chat_display = scrolledtext.ScrolledText(right_frame,
                                                      wrap=tk.WORD,
                                                      font=('Segoe UI', 9),
                                                      state=tk.DISABLED,
                                                      bg='#f8f9fa',
                                                      padx=10,
                                                      pady=10)
        self.chat_display.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure text tags
        self.chat_display.tag_configure('user', foreground='#2980b9', font=('Segoe UI', 9, 'bold'))
        self.chat_display.tag_configure('assistant', foreground='#27ae60', font=('Segoe UI', 9, 'bold'))
        self.chat_display.tag_configure('system', foreground='#e74c3c', font=('Segoe UI', 9, 'italic'))
        self.chat_display.tag_configure('timestamp', foreground='#95a5a6', font=('Segoe UI', 8))
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=(15, 0))
        
        ttk.Button(button_frame, text="Load Into Chat", command=self.load_into_current).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="🗑️ Delete", command=self.delete_chat).grid(row=0, column=1, padx=5)
        ttk.Button(button_frame, text="📁 Open Folder", command=self.open_folder).grid(row=0, column=2, padx=5)
        ttk.Button(button_frame, text="Close", command=self.destroy).grid(row=0, column=3, padx=5)
        
        # Load chats
        self.load_chat_list()
    
    def load_chat_list(self):
        """Load list of saved chat files"""
        self.chat_listbox.delete(0, tk.END)
        self.chat_files = []
        
        try:
            # Get all chat JSON files, sorted newest first
            files = sorted(self.app.chat_history_dir.glob("chat_*.json"),
                          key=lambda f: f.stat().st_mtime,
                          reverse=True)
            
            for file in files:
                self.chat_files.append(file)
                
                # Format: chat_YYYYMMDD_HHMMSS.json
                try:
                    filename = file.stem
                    timestamp = filename.replace('chat_', '')
                    date_part = timestamp[:8]  # YYYYMMDD
                    time_part = timestamp[9:]  # HHMMSS
                    
                    display = f"{date_part[:4]}-{date_part[4:6]}-{date_part[6:8]} {time_part[:2]}:{time_part[2:4]}:{time_part[4:6]}"
                except:
                    display = file.name
                
                self.chat_listbox.insert(tk.END, display)
            
            if not files:
                self.info_label.config(text="No saved chats found")
        
        except Exception as e:
            logger.error(f"Error loading chat list: {e}")
            self.info_label.config(text="Error loading chat list")
    
    def on_chat_select(self, event):
        """Display selected chat"""
        selection = self.chat_listbox.curselection()
        if not selection:
            return
        
        index = selection[0]
        file = self.chat_files[index]
        
        try:
            with open(file, 'r') as f:
                data = json.load(f)
            
            # Update info
            saved_at = data.get('saved_at', '')
            messages = data.get('messages', [])
            
            try:
                dt = datetime.fromisoformat(saved_at)
                date_str = dt.strftime("%B %d, %Y at %I:%M %p")
            except:
                date_str = saved_at
            
            self.info_label.config(text=f"Saved: {date_str} | {len(messages)} messages")
            
            # Display messages
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete(1.0, tk.END)
            
            for msg in messages:
                timestamp = msg.get('timestamp', '')
                sender = msg.get('sender', 'Unknown')
                text = msg.get('message', '')
                msg_type = msg.get('type', 'user')
                
                # Format timestamp
                try:
                    ts = datetime.fromisoformat(timestamp)
                    ts_str = ts.strftime("%H:%M")
                except:
                    ts_str = timestamp[:5] if timestamp else ""
                
                self.chat_display.insert(tk.END, f"[{ts_str}] ", 'timestamp')
                self.chat_display.insert(tk.END, f"{sender}: ", msg_type)
                self.chat_display.insert(tk.END, f"{text}\n\n")
            
            self.chat_display.config(state=tk.DISABLED)
            self.chat_display.see(1.0)
        
        except Exception as e:
            logger.error(f"Error loading chat: {e}")
            messagebox.showerror("Error", f"Failed to load chat: {e}", parent=self)
    
    def load_into_current(self):
        """Load selected chat into current conversation"""
        selection = self.chat_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a chat to load", parent=self)
            return
        
        index = selection[0]
        file = self.chat_files[index]
        
        if messagebox.askyesno("Load Chat",
                              "This will replace your current conversation.\n\n"
                              "Load this chat into the main window?",
                              parent=self):
            try:
                with open(file, 'r') as f:
                    data = json.load(f)
                
                # Clear current chat
                self.app.chat_display.config(state=tk.NORMAL)
                self.app.chat_display.delete(1.0, tk.END)
                self.app.chat_display.config(state=tk.DISABLED)
                
                # Load messages
                messages = data.get('messages', [])
                self.app.current_chat_history = messages.copy()
                
                # Display in main window
                for msg in messages:
                    sender = msg.get('sender', 'Unknown')
                    text = msg.get('message', '')
                    msg_type = msg.get('type', 'user')
                    
                    self.app.add_message(sender, text, msg_type)
                
                messagebox.showinfo("Chat Loaded", f"Loaded {len(messages)} messages", parent=self)
                self.destroy()
            
            except Exception as e:
                logger.error(f"Error loading chat: {e}")
                messagebox.showerror("Error", f"Failed to load chat: {e}", parent=self)
    
    def delete_chat(self):
        """Delete selected chat file"""
        selection = self.chat_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a chat to delete", parent=self)
            return
        
        index = selection[0]
        file = self.chat_files[index]
        
        if messagebox.askyesno("Delete Chat",
                              f"Permanently delete this chat?\n\n{file.name}",
                              parent=self):
            try:
                file.unlink()
                
                # Clear display
                self.chat_display.config(state=tk.NORMAL)
                self.chat_display.delete(1.0, tk.END)
                self.chat_display.config(state=tk.DISABLED)
                self.info_label.config(text="Select a chat to view")
                
                # Reload list
                self.load_chat_list()
                
                logger.info(f"Deleted chat: {file.name}")
            
            except Exception as e:
                logger.error(f"Error deleting chat: {e}")
                messagebox.showerror("Error", f"Failed to delete chat: {e}", parent=self)
    
    def open_folder(self):
        """Open chat history folder in file explorer"""
        import subprocess
        import sys
        
        try:
            path = str(self.app.chat_history_dir)
            
            if sys.platform == 'win32':
                subprocess.run(['explorer', path])
            elif sys.platform == 'darwin':
                subprocess.run(['open', path])
            else:
                subprocess.run(['xdg-open', path])
        
        except Exception as e:
            logger.error(f"Error opening folder: {e}")
            messagebox.showerror("Error", f"Failed to open folder: {e}", parent=self)


def main():
    """Main entry point"""
    print("\n" + "="*60)
    print("Garmin Chat - Desktop Application")
    print("="*60)
    print("\nStarting application...")
    print("="*60 + "\n")
    
    root = tk.Tk()
    app = GarminChatApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
